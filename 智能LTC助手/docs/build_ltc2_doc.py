from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "LTC2.0智能销售作战与资源加速系统设计.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("LTC 2.0 智能销售作战与资源加速系统设计")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.font.name = "Calibri"
    set_east_asia_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("系统架构设计、功能设计与功能说明")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")
    set_east_asia_font(run)


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F2F4F7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
    for cell in table.rows[0].cells:
        set_cell_margins(cell)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


doc = Document()
style_doc(doc)
add_title(doc)

add_para(doc, "本文档基于本次讨论整理，用于后续产品原型、需求分析与程序开发。系统不是传统 CRM，也不是传统 LTC 流程固化工具，而是一套以销售为中心、由实时信息驱动的动态销售作战系统。")

doc.add_heading("一、系统定位", level=1)
add_para(doc, "系统名称暂定为：LTC 2.0 智能销售作战与资源加速系统。")
add_para(doc, "核心定位：销售实时输入客户信息，系统自动理解商机状态，动态推荐下一步动作，帮助销售争取公司资源、加快审批流程、提升赢单概率，并为管理者提供更真实的销售预测。")
add_para(doc, "价值主张：销售提供真实信息，系统帮助销售换取资源、加速流程、打赢商机。")

doc.add_heading("二、核心理念", level=1)
add_table(doc, ["传统 LTC", "LTC 2.0"], [
    ["先固化流程，再要求销售执行", "先理解销售现场，再动态生成任务、资源和流程动作"],
    ["固定阶段", "动态商机状态"],
    ["销售填字段", "销售说人话"],
    ["流程管控", "资源加速"],
    ["固定审批", "风险触发审批"],
    ["管理者追问", "系统自动摘要"],
    ["销售被监督", "销售被赋能"],
    ["靠经验预测", "基于状态预测"],
], [3.1, 3.4])

doc.add_heading("三、总体架构", level=1)
add_table(doc, ["架构层", "核心能力", "典型输入/输出"], [
    ["输入层", "销售日志、会议纪要、客户互动记录、CRM/报价/合同数据接入", "销售自然语言、语音转文字、邮件/会议内容、业务系统数据"],
    ["AI 理解层", "信息抽取、风险识别、关键人识别、下一步动作识别", "结构化客户需求、预算、时间、竞争、风险、关键人、待办"],
    ["商机状态层", "需求清晰度、预算确定性、决策链清晰度、竞争风险、商务风险、交付风险等状态画像", "商机成熟度、风险等级、信息完整度、资源权益等级"],
    ["决策与推荐层", "任务推荐、资源推荐、快速审批判断、赢率预测、风险预警", "下一步动作、资源申请建议、审批路径、预测结果"],
    ["业务应用层", "销售作战台、商机智能详情页、资源加速器、管理者驾驶舱、审批与协同中心", "面向销售、主管、售前、交付、财务、法务和管理层的工作界面"],
], [1.25, 2.65, 2.6])

doc.add_heading("四、核心业务闭环", level=1)
add_numbered(doc, [
    "销售输入真实信息。",
    "系统理解客户现场。",
    "生成商机状态画像。",
    "推荐下一步动作。",
    "解锁公司资源。",
    "加快审批与协同。",
    "提升赢单概率。",
    "销售更愿意持续使用系统，形成正向循环。",
])

doc.add_heading("五、用户角色设计", level=1)
add_table(doc, ["角色", "主要诉求", "系统价值"], [
    ["销售 AR", "快拿资源、快过审批、快赢单", "智能建议、资源申请、自动材料、流程加速"],
    ["售前 SR", "判断哪些商机值得支持", "商机质量评分、需求摘要、支持优先级"],
    ["交付 FR", "避免销售乱承诺", "交付风险识别、提前介入、交底信息完整"],
    ["销售主管", "看真实商机、少靠口头汇报", "商机可信度、风险预警、预测修正"],
    ["财务/法务", "控风险、提效率", "风险触发审批、自动生成审批材料"],
    ["管理层", "看预测、看资源投入产出", "经营驾驶舱、重点商机、资源瓶颈"],
], [1.2, 2.2, 3.1])

doc.add_heading("六、核心模块设计", level=1)

doc.add_heading("1. 销售实时输入模块", level=2)
add_para(doc, "功能定位：降低销售录入成本，让销售可以像写拜访记录一样输入客户信息。")
add_bullets(doc, ["文本输入", "语音转文字", "会议纪要导入", "邮件/聊天记录复制", "客户拜访记录", "电话沟通记录", "手动补充字段"])
add_para(doc, "示例输入：今天和客户信息部张总、采购王经理开会。张总认可我们的方案，但采购认为价格偏高。客户预算大约 80 万，希望 8 月完成选型。目前 A 厂商也在跟进，报价可能更低。客户希望我们下周提供行业案例和初步报价。")
add_table(doc, ["信息项", "识别结果"], [
    ["关键人", "信息部张总、采购王经理"],
    ["客户态度", "信息部认可方案"],
    ["风险", "采购压价、竞品进入"],
    ["预算", "约 80 万"],
    ["时间", "8 月完成选型"],
    ["竞争对手", "A 厂商"],
    ["下一步", "提供行业案例、初步报价"],
], [1.6, 4.9])
add_bullets(doc, ["自动提取客户需求", "自动识别关键人、预算、时间、竞争和风险", "自动生成下一步待办", "自动更新商机状态", "自动提示缺失信息"])

doc.add_heading("2. 商机状态画像模块", level=2)
add_para(doc, "功能定位：不再只用销售阶段描述商机，而是用多个状态维度描述商机真实成熟度。")
add_table(doc, ["状态维度", "说明"], [
    ["需求清晰度", "是否明确客户真实业务问题"],
    ["痛点强度", "客户问题是否足够迫切"],
    ["预算确定性", "是否有预算、预算是否审批"],
    ["决策链清晰度", "是否知道谁影响、谁拍板"],
    ["关键人支持度", "是否有内部支持者"],
    ["采购时间明确度", "是否有明确采购计划"],
    ["竞争风险", "是否有竞品、是否价格战"],
    ["方案匹配度", "方案是否解决客户核心问题"],
    ["商务风险", "折扣、付款、合同条款风险"],
    ["交付风险", "周期、定制、资源、验收风险"],
    ["回款风险", "付款条件、信用、历史回款情况"],
    ["信息完整度", "商机信息是否足够支撑判断"],
], [1.8, 4.7])
add_para(doc, "输出结果包括：当前成熟度、风险等级、赢率预测、资源建议、缺失信息和下一步动作。")

doc.add_heading("3. 动态任务推荐模块", level=2)
add_para(doc, "功能定位：系统根据商机状态自动推荐任务，而不是预先固定所有流程任务。")
add_table(doc, ["系统识别状态", "推荐任务"], [
    ["预算不清", "确认预算来源、预算上限、审批路径"],
    ["决策人缺失", "补充决策链地图，识别最终拍板人"],
    ["竞品进入", "准备竞品对比、差异化价值材料"],
    ["客户要求报价过早", "先完成需求澄清和预算确认"],
    ["非标交付", "申请交付评估"],
    ["折扣过高", "生成折扣审批材料"],
    ["商机停滞", "提醒销售重新确认客户采购计划"],
    ["付款条件异常", "触发财务风险评审"],
], [2.0, 4.5])

doc.add_heading("4. 销售资源加速器模块", level=2)
add_para(doc, "功能定位：让销售感觉系统是在帮他争取资源，而不是监督他。这是系统最重要的销售侧价值模块。")
add_table(doc, ["功能", "说明"], [
    ["商机权益等级", "根据信息完整度和商机质量解锁资源"],
    ["一键资源申请", "申请售前、交付、法务、财务、高层支持"],
    ["自动申请材料", "系统根据商机信息自动生成申请理由"],
    ["资源优先级评估", "判断商机是否值得优先投入资源"],
    ["资源 SLA 跟踪", "显示资源响应人、响应时间、处理状态"],
    ["升级提醒", "超时后可提醒主管或升级处理"],
], [1.8, 4.7])
add_table(doc, ["等级", "条件", "可获得权益"], [
    ["L0 信息不足", "只有客户名和金额", "仅记录商机"],
    ["L1 基础有效", "有客户需求和联系人", "可进入销售漏斗"],
    ["L2 需求验证", "有明确痛点和场景", "可申请售前初步支持"],
    ["L3 商机明确", "有预算、时间、需求", "可申请方案资源"],
    ["L4 决策清晰", "有决策链和关键人支持", "可申请主管陪访、高层支持"],
    ["L5 商务成熟", "报价、交付、付款风险可控", "可进入快速审批"],
    ["L6 签约冲刺", "客户决策时间明确", "可申请跨部门加急支持"],
], [1.4, 2.3, 2.8])

doc.add_heading("5. 快速审批通道模块", level=2)
add_para(doc, "功能定位：让信息充分、风险可控的商机更快通过报价、折扣、合同等流程。")
add_bullets(doc, ["客户预算明确", "需求已确认", "决策链基本清楚", "折扣在授权范围内", "毛利达标", "无重大非标条款", "交付资源已确认", "付款条件符合政策"])
add_para(doc, "如果符合条件，系统提示：当前商机符合快速报价审批条件，可进入快速通道。")
add_para(doc, "如果不符合条件，系统提示：当前缺少 2 项信息，如预算来源、交付周期确认。补充后可进入快速审批判断。")

doc.add_heading("6. 风险预警模块", level=2)
add_para(doc, "功能定位：从销售输入和商机状态中识别风险，并给出处理建议。")
add_table(doc, ["风险", "说明", "建议动作"], [
    ["预算风险", "客户预算不明或低于报价", "确认预算来源，设计分阶段方案"],
    ["决策风险", "未识别最终决策人", "建立决策链地图"],
    ["竞争风险", "竞品进入或低价竞争", "准备差异化价值材料"],
    ["报价风险", "客户过早要求报价", "先完成需求和预算澄清"],
    ["交付风险", "非标定制或周期过短", "申请交付评估"],
    ["合同风险", "特殊付款或违约条款", "触发财务/法务评审"],
    ["停滞风险", "商机长期无进展", "降低赢率，提醒复盘"],
    ["回款风险", "付款条件差或客户信用弱", "财务提前介入"],
], [1.3, 2.5, 2.7])

doc.add_heading("7. 智能销售教练模块", level=2)
add_para(doc, "功能定位：帮助销售更好推进客户，而不只是记录客户信息。")
add_bullets(doc, ["异议应对建议", "价格谈判建议", "竞品对比建议", "ROI 沟通建议", "决策链推进建议", "下一次拜访问题清单", "高层汇报话术", "客户行业案例推荐", "折扣交换条件建议"])
add_para(doc, "示例：客户觉得我们比 A 厂商贵 20%。系统建议销售不要立即降价，而是先确认竞品报价是否包含实施和服务，引导客户比较总拥有成本，准备 ROI 测算，并在必须降价时换取预付款比例或签约时间承诺。")

doc.add_heading("8. 自动材料生成模块", level=2)
add_para(doc, "功能定位：减少销售重复写材料，提高申请和审批效率。")
add_bullets(doc, ["售前支持申请", "折扣申请说明", "高层拜访背景", "客户需求摘要", "方案交底材料", "合同特批说明", "交付交底单", "回款风险说明", "商机复盘报告", "输单分析报告", "周报/月报摘要"])
add_para(doc, "关键价值：销售认真录一次，系统自动复用多次。")

doc.add_heading("9. 管理者驾驶舱模块", level=2)
add_para(doc, "功能定位：让管理者看到真实商机质量，而不是只看销售主观填报。")
add_table(doc, ["看板", "内容"], [
    ["销售预测看板", "预测金额、赢率、签约时间、可信度"],
    ["高风险商机看板", "大额、高竞争、高折扣、交付风险商机"],
    ["停滞商机看板", "长时间无更新、无下一步动作的商机"],
    ["资源投入看板", "售前、交付、法务、财务资源使用情况"],
    ["销售过程质量看板", "信息完整度、任务完成率、预测偏差"],
    ["审批效率看板", "各类审批平均耗时、超时事项"],
    ["漏斗健康度看板", "阶段转化率、商机流失原因、赢单率"],
], [1.8, 4.7])

doc.add_heading("七、关键数据对象设计", level=1)
add_table(doc, ["对象", "说明"], [
    ["Customer 客户", "客户基本信息、行业、规模、等级"],
    ["Contact 联系人", "客户联系人、角色、影响力、态度"],
    ["Opportunity 商机", "金额、预计签约时间、状态、赢率"],
    ["Interaction 互动记录", "拜访、电话、会议、邮件、纪要"],
    ["Signal 业务信号", "从互动中识别出的预算、风险、竞争等"],
    ["Task 任务", "系统推荐或人工创建的下一步动作"],
    ["Risk 风险", "风险类型、等级、原因、处理建议"],
    ["ResourceRequest 资源申请", "售前、交付、法务、主管等资源请求"],
    ["Approval 审批", "折扣、合同、非标、交付等审批"],
    ["Forecast 预测", "赢率、金额、时间、可信度"],
    ["Material 材料", "系统自动生成的申请、交底、汇报材料"],
], [2.0, 4.5])

doc.add_heading("八、MVP 版本建议", level=1)
add_para(doc, "第一版不要做太大，建议先做 6 个核心能力。")
add_numbered(doc, [
    "销售自然语言日志输入。",
    "AI 自动提取客户信息和风险。",
    "商机状态评分。",
    "下一步任务推荐。",
    "一键资源申请和自动生成申请理由。",
    "管理者基础驾驶舱。",
])
add_table(doc, ["页面", "作用"], [
    ["销售个人作战台", "销售每天看今天该推进什么"],
    ["商机智能详情页", "查看商机状态、风险、任务、资源"],
    ["管理者驾驶舱", "查看预测、高风险商机、资源请求"],
], [2.0, 4.5])

doc.add_heading("九、推荐页面设计", level=1)
doc.add_heading("1. 销售个人作战台", level=2)
add_bullets(doc, ["今天最该推进的商机", "每个商机的下一步建议", "当前可解锁的资源", "需要补充的信息", "审批和资源申请进度", "高风险提醒", "本周预计签约机会"])
add_para(doc, "核心体验：销售不用找系统，系统告诉销售今天该干什么。")

doc.add_heading("2. 商机智能详情页", level=2)
add_bullets(doc, ["商机摘要", "当前状态评分", "关键信息缺口", "系统推荐动作", "风险预警", "资源权益等级", "一键资源申请", "客户关键人地图", "互动记录时间线", "审批与协同记录"])

doc.add_heading("3. 资源加速中心", level=2)
add_para(doc, "销售可以看到：我申请了什么资源、谁在处理、预计什么时候完成、是否超时、是否可以升级、处理结果是什么。资源方可以看到：哪些商机需要支持、商机质量如何、支持优先级如何、销售提供的信息是否充分。")

doc.add_heading("4. 管理者驾驶舱", level=2)
add_bullets(doc, ["本月/本季度预测金额", "预测可信度", "高风险大商机", "停滞商机", "资源瓶颈", "审批超时", "销售过程质量", "需要管理者介入的商机"])

doc.add_heading("十、技术架构建议", level=1)
add_table(doc, ["技术层", "建议能力"], [
    ["前端", "Web 管理端、移动端 H5 或小程序，销售移动录入优先"],
    ["后端业务服务", "用户与权限、商机管理、互动记录管理、任务管理、资源申请、审批流、看板统计"],
    ["AI 服务", "自然语言信息抽取、风险识别、商机摘要生成、推荐动作生成、材料自动生成、销售教练问答"],
    ["规则引擎", "商机评分规则、风险触发规则、资源权益规则、快速审批规则、赢率预测规则"],
    ["数据层", "结构化业务数据库、非结构化互动记录库、向量知识库"],
    ["集成层", "CRM、ERP、合同系统、企业微信/钉钉、邮件和会议系统"],
], [1.6, 4.9])

doc.add_heading("十一、系统成功的关键指标", level=1)
add_table(doc, ["指标", "说明"], [
    ["销售日志录入率", "销售是否愿意持续输入"],
    ["信息完整度提升", "商机关键字段是否更完整"],
    ["资源申请响应时间", "售前、交付、法务响应是否更快"],
    ["审批周期缩短", "报价、折扣、合同审批是否提速"],
    ["商机预测准确率", "系统预测是否接近真实结果"],
    ["赢单率提升", "重点商机转化是否改善"],
    ["销售重复材料减少", "是否减少申请、汇报、交底工作"],
    ["管理追问减少", "主管是否减少重复询问"],
    ["资源投入产出比", "高价值商机是否获得更多支持"],
], [2.0, 4.5])

doc.add_heading("十二、一句话总结", level=1)
add_para(doc, "这套 LTC 2.0 系统的本质不是流程系统，而是一个以销售实时信息为燃料、以商机状态为核心、以资源加速为激励、以动态任务和风险判断为能力的智能销售作战系统。")
add_para(doc, "后续开发建议第一步围绕销售日志输入、商机状态画像、资源加速器、动态任务推荐这四个模块做 MVP。这个方向最容易让销售感受到价值，也最能区别于传统 CRM/LTC 系统。")

doc.core_properties.title = "LTC 2.0 智能销售作战与资源加速系统设计"
doc.core_properties.subject = "系统架构设计、功能设计与功能说明"
doc.core_properties.author = "Codex"

doc.save(OUT)
print(OUT)
